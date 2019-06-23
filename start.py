from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag

import urllib2, StringIO

import requests

import os
from datetime import datetime

curpath = os.path.abspath(os.curdir)

from docx import Document
from docx.shared import Inches

from urlparse import urljoin

BASE_LINK = "http://www.keyakizaka46.com"
LINK = "http://www.keyakizaka46.com/s/k46o/diary/member/list?ima=0000&page=%s&cd=member&ct=%s"

first_page = 0
last_page = 0
idol = "10"

from fpdf import FPDF, HTMLMixin

class BlogEntry(object):
    title = ""
    link = ""
    date = datetime.now()
    contents = []

class BlogContent(object):
    def get_content(self):
        pass

    def change_document(self, document):
        pass

class BlogHeadingContent(BlogContent):
    def __init__(self, heading, level=1):
        self.heading = heading
        self.level = level
    
    def get_content(self):
        return self.heading
    
    def change_document(self, document):
        document.add_heading(self.heading, level=self.level)

class BlogTextContent(BlogContent):
    def __init__(self, text):
        self.text = text

    def get_content(self):
        return self.text

    def change_document(self, document):
        document.add_paragraph(self.text)

class BlogImageContent(BlogContent):
    def __init__(self, img_url):
        self.img_url = img_url

    def get_content(self):
        image_from_url = urllib2.urlopen(self.img_url)
        io_url = StringIO.StringIO()
        io_url.write(image_from_url.read())
        io_url.seek(0)
        return io_url

    def change_document(self, document):
        try:
            document.add_picture(self.get_content())
        except:
            pass

def traverse_through_article(element):
    contents = []

    children = element.children
    for child in children:
        # print child
        if type(child) is NavigableString:
            # print child
            contents.append(BlogTextContent(child))
        elif type(child) is Tag:
            if child.name == 'p':
                contents.append(BlogTextContent(child.get_text()))
                # print child.get_text()
            elif child.name == 'img':
                # print child.get('src')
                contents.append(BlogImageContent(child.get('src')))
            elif child.name == 'a':
                    contents.append(BlogTextContent(child.get_text()))
            elif child.name == 'br':
                contents.append(BlogTextContent(''))
            elif child.name == 'div' or child.name == 'span':
                contents.extend(traverse_through_article(child))
    return contents

for page_number in range(first_page, last_page + 1):
    pageLink = LINK % (page_number, idol)
    print pageLink

    req = requests.get(pageLink)
    soup = BeautifulSoup(req.text, "lxml")
    blog_lists = soup.find_all("div", class_="keyaki-blog_list")
    for blog_list in blog_lists:
        articles = blog_list.find_all('article')
        for article in articles:
            metaContainer = article.find("div", class_="innerHead")
            bottomContainer = article.find("div", class_="box-bottom")
            bottomContainerUl = bottomContainer.find("ul")
            bottomContainerLi = bottomContainerUl.find_all("li")

            blogTitleContainer = metaContainer.find("div", class_="box-ttl")
            blogTitle = blogTitleContainer.find("h3")
            blogTitleAnchor = blogTitleContainer.find("a")
            blogTitleLink = urljoin(BASE_LINK, blogTitleAnchor.get("href"))
            blogFileName = "%s - %s" % (bottomContainerLi[0].get_text().strip(), blogTitle.get_text().strip())
            print blogFileName
            blogFile = os.path.join(curpath, 'output', '%s' % (blogFileName.replace("/","-").replace("\\","-")))

            contents = []
            contents.append(BlogHeadingContent(blogFileName))
            contents.append(BlogHeadingContent(blogTitleLink, 2))
            # document.add_heading(blogFileName, level=1)
            # document.add_heading(blogTitleLink, level=2)
            contents.extend(traverse_through_article(article.find("div", class_="box-article")))
            print "===================="
            document = Document()

            style = document.styles['Normal']
            headingStyle = document.styles['Heading 1']
            headingLinkStyle = document.styles['Heading 2']
            font = style.font
            headingFont = headingStyle.font
            font.name = "Hiragino Sans"
            headingFont.name = "Hiragino Sans"
            headingLinkStyle.font.name = "Hiragino Sans"

            for content in contents:
                content.change_document(document)
            document.save("%s.docx" % (blogFile))
    
    